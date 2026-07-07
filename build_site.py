from pathlib import Path
from html import escape
import docx
from urllib.parse import quote

DOCX = Path('/mnt/data/אחרון ודי.docx')
SITE = Path('/mnt/data/aspect-hebrew-site')

doc = docx.Document(DOCX)
paras = {i: p.text.strip() for i, p in enumerate(doc.paragraphs)}
final_callout = doc.tables[0].cell(0, 0).text.strip() if doc.tables else ''


def p_html(text, cls=''):
    cls_attr = f' class="{cls}"' if cls else ''
    return f'<p{cls_attr}>{escape(text)}</p>'


def bullet_html(text):
    text = text.strip()
    if text.startswith('●'):
        text = text[1:].strip()
    # Preserve the document's natural emphasis by bolding the first sentence of each bullet.
    if '.  ' in text:
        title, rest = text.split('.  ', 1)
        return f'<li><strong>{escape(title)}.</strong> {escape(rest)}</li>'
    if '. ' in text:
        title, rest = text.split('. ', 1)
        return f'<li><strong>{escape(title)}.</strong> {escape(rest)}</li>'
    return f'<li>{escape(text)}</li>'


def paragraphs(indices, lead_first=False):
    out = []
    for n, idx in enumerate(indices):
        text = paras.get(idx, '')
        if not text or text == '.':
            continue
        out.append(p_html(text, 'lead' if lead_first and n == 0 else ''))
    return '\n'.join(out)


def section(sec_id, eyebrow, title, body, image=None, image_alt='', variant=''):
    media = ''
    if image:
        media = f'''
        <div class="section-media">
          <img src="assets/{image}" alt="{escape(image_alt)}" loading="lazy">
        </div>'''
    return f'''
    <section class="content-section {variant}" id="{sec_id}">
      <div class="section-copy">
        <div class="eyebrow">{escape(eyebrow)}</div>
        <h2>{escape(title)}</h2>
        {body}
      </div>
      {media}
    </section>'''

intro = paragraphs(range(3, 14), lead_first=True)
learn = paragraphs(range(17, 21))
houston = paragraphs(range(22, 24))
why_houston = paragraphs([26]) + '\n<ul class="feature-list">' + '\n'.join(bullet_html(paras[i]) for i in range(27, 33)) + '</ul>'
why_build = paragraphs([34]) + '\n<ul class="feature-list two-col">' + '\n'.join(bullet_html(paras[i]) for i in range(35, 44)) + '</ul>'
why_now = paragraphs(range(45, 50))
private_asset = paragraphs(range(51, 56))
role = paragraphs(range(58, 64))
capital = paragraphs(range(65, 67))
time_strategy = paragraphs(range(68, 72))
end_to_end = paragraphs(range(73, 80))
long_hold = paragraphs(range(84, 88))
next_steps = paragraphs(range(89, 92))
callout_html = ''.join(p_html(t.strip()) for t in final_callout.split('\n') if t.strip())
whatsapp_href = 'https://wa.me/972537650555?text=' + quote('שלום, הגעתי מאתר אספקט ואשמח לשמוע עוד על מודל ההשקעה ביוסטון.')

html = f'''<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1, viewport-fit=cover">
  <meta name="theme-color" content="#143a46">
  <meta name="description" content="Aspect Boutique Investments - רכישת קרקעות, ייזום ובניית בתים פרטיים בשכונות מתפתחות ביוסטון, טקסס.">
  <title>{escape(paras[1])} | Aspect Boutique Investments</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Assistant:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
  <style>
    :root {{
      --teal:#143a46;
      --teal-2:#1d4e5d;
      --blue:#29ABE2;
      --blue-soft:#eaf6fc;
      --ink:#17262d;
      --body:#33444d;
      --muted:#71828b;
      --line:#dfe8ec;
      --surface:#ffffff;
      --tint:#f6f9fb;
      --sand:#f5f1ea;
      --max:1180px;
      --text:760px;
      --shadow:0 24px 70px -38px rgba(20,58,70,.55);
      --radius:26px;
    }}
    *{{box-sizing:border-box;margin:0;padding:0}}
    html{{scroll-behavior:smooth}}
    body{{
      font-family:'Assistant',system-ui,-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
      background:var(--surface);color:var(--body);font-size:17px;line-height:1.78;
      -webkit-text-size-adjust:100%;text-rendering:optimizeLegibility;
    }}
    img{{max-width:100%;display:block}}
    a{{color:inherit}}
    .progress{{position:fixed;top:0;right:0;height:3px;background:var(--blue);width:0;z-index:100}}
    .topbar{{position:sticky;top:0;z-index:80;background:rgba(255,255,255,.94);backdrop-filter:blur(14px);border-bottom:1px solid rgba(223,232,236,.78)}}
    .topbar .inner{{max-width:var(--max);margin:auto;display:flex;align-items:center;justify-content:space-between;gap:20px;padding:10px 22px}}
    .brand{{display:flex;align-items:center;text-decoration:none;min-width:176px}}
    .brand-logo{{width:230px;height:60px;object-fit:cover;object-position:50% 50%;border-radius:0}}
    nav{{display:flex;align-items:center;gap:22px;font-size:15.5px;color:var(--teal);font-weight:700}}
    nav a{{text-decoration:none;opacity:.86;transition:opacity .16s,color .16s}}
    nav a:hover{{opacity:1;color:var(--blue)}}
    .nav-cta{{background:var(--teal);color:#fff!important;padding:9px 15px;border-radius:999px;opacity:1}}

    .hero{{position:relative;min-height:76svh;display:grid;align-items:end;overflow:hidden;background:var(--teal);isolation:isolate}}
    .hero::before{{content:"";position:absolute;inset:0;background:url('assets/homes-exterior.jpg') center/cover no-repeat;z-index:-3;transform:scale(1.04)}}
    .hero::after{{content:"";position:absolute;inset:0;background:
      radial-gradient(circle at 20% 20%,rgba(41,171,226,.35),transparent 34%),
      linear-gradient(90deg,rgba(6,25,31,.20),rgba(6,25,31,.86) 52%,rgba(6,25,31,.96));z-index:-2}}
    .hero-inner{{width:100%;max-width:var(--max);margin:auto;padding:84px 24px 58px;display:grid;grid-template-columns:minmax(0,650px) minmax(250px,410px);gap:54px;align-items:end}}
    .hero-card{{color:#fff;padding:34px;border:1px solid rgba(255,255,255,.18);background:rgba(9,36,45,.42);box-shadow:0 30px 80px -44px #000;border-radius:30px;backdrop-filter:blur(10px)}}
    .hero-badge{{display:inline-flex;align-items:center;gap:8px;background:rgba(41,171,226,.16);border:1px solid rgba(41,171,226,.45);color:#dff6ff;padding:6px 13px;border-radius:999px;font-size:14px;font-weight:700;margin-bottom:18px}}
    .hero h1{{font-size:clamp(32px,5vw,62px);line-height:1.08;font-weight:800;letter-spacing:-.02em;color:#fff;margin-bottom:18px}}
    .hero .subtitle{{font-size:clamp(20px,2.4vw,29px);font-weight:700;color:#dff6ff;margin-bottom:22px;line-height:1.25}}
    .hero .summary{{font-size:19px;color:#d8e8ed;max-width:610px;margin-bottom:28px}}
    .hero-actions{{display:flex;flex-wrap:wrap;gap:13px}}
    .btn{{display:inline-flex;align-items:center;justify-content:center;gap:8px;border-radius:14px;padding:13px 20px;text-decoration:none;font-weight:800;line-height:1.2;transition:transform .16s,box-shadow .16s,background .16s}}
    .btn:hover{{transform:translateY(-1px)}}
    .btn.primary{{background:#fff;color:var(--teal);box-shadow:0 20px 35px -24px #000}}
    .btn.secondary{{background:rgba(255,255,255,.08);color:#fff;border:1px solid rgba(255,255,255,.34)}}
    .hero-panel{{background:rgba(255,255,255,.95);border-radius:30px;padding:25px;box-shadow:var(--shadow);border:1px solid rgba(255,255,255,.64)}}
    .hero-panel .logo-box{{background:#030303;border-radius:22px;padding:16px 18px 12px;margin-bottom:18px}}
    .hero-panel img.logo-large{{width:100%;height:86px;object-fit:cover;object-position:50% 50%;margin:auto}}
    .metrics{{display:grid;gap:12px}}
    .metric{{padding:14px 15px;border-radius:17px;background:var(--tint);border:1px solid var(--line)}}
    .metric strong{{display:block;color:var(--teal);font-size:20px;line-height:1.1;margin-bottom:3px}}
    .metric span{{color:var(--muted);font-size:14.5px}}

    main{{max-width:var(--max);margin:auto;padding:38px 24px 0}}
    .intro{{display:grid;grid-template-columns:minmax(0,1fr) 350px;gap:42px;align-items:start;padding:42px 0 22px}}
    .intro-copy{{max-width:var(--text)}}
    .intro-copy .lead{{font-size:20px;color:var(--ink);font-weight:600;line-height:1.74}}
    .intro p,.content-section p{{margin:0 0 15px}}
    .intro p:last-child,.content-section p:last-child{{margin-bottom:0}}
    .aside-card{{position:sticky;top:92px;background:linear-gradient(180deg,var(--tint),#fff);border:1px solid var(--line);border-radius:var(--radius);padding:24px;box-shadow:var(--shadow)}}
    .aside-card h3{{font-size:23px;color:var(--teal);line-height:1.2;margin-bottom:13px}}
    .aside-card p{{font-size:16px;color:var(--body)}}
    .aside-card .mini{{display:grid;gap:10px;margin-top:18px}}
    .aside-card .mini div{{display:flex;gap:10px;align-items:flex-start;background:#fff;border:1px solid var(--line);border-radius:15px;padding:12px}}
    .aside-card .mini div::before{{content:"";flex:0 0 9px;width:9px;height:9px;border-radius:999px;background:var(--blue);margin-top:9px}}

    .content-section{{display:grid;grid-template-columns:minmax(0,1fr) 430px;gap:46px;align-items:center;padding:74px 0;border-top:1px solid var(--line)}}
    .content-section.wide{{grid-template-columns:minmax(0,1fr)}}
    .content-section.reverse{{grid-template-columns:430px minmax(0,1fr)}}
    .content-section.reverse .section-copy{{grid-column:2}}
    .content-section.reverse .section-media{{grid-column:1;grid-row:1}}
    .section-copy{{max-width:var(--text)}}
    .eyebrow{{display:inline-flex;color:var(--blue);font-size:14px;font-weight:800;letter-spacing:.02em;margin-bottom:8px}}
    h2{{font-size:clamp(27px,3.2vw,42px);line-height:1.18;color:var(--teal);font-weight:800;letter-spacing:-.02em;margin-bottom:20px;position:relative;padding-bottom:14px}}
    h2::after{{content:"";position:absolute;right:0;bottom:0;width:64px;height:4px;border-radius:999px;background:var(--blue)}}
    .section-media{{border-radius:30px;overflow:hidden;box-shadow:var(--shadow);background:var(--tint);min-height:330px}}
    .section-media img{{width:100%;height:100%;min-height:330px;object-fit:cover}}
    .feature-list{{list-style:none;display:grid;gap:13px;margin-top:18px}}
    .feature-list.two-col{{grid-template-columns:repeat(2,minmax(0,1fr))}}
    .feature-list li{{position:relative;background:var(--tint);border:1px solid var(--line);border-radius:18px;padding:16px 44px 16px 18px;font-size:16.5px;line-height:1.65;color:var(--ink)}}
    .feature-list li::before{{content:"";position:absolute;right:15px;top:21px;width:16px;height:16px;border-radius:6px;background:var(--blue-soft)}}
    .feature-list li::after{{content:"";position:absolute;right:19px;top:24px;width:8px;height:5px;border-right:2px solid var(--blue);border-bottom:2px solid var(--blue);transform:rotate(45deg)}}
    .feature-list strong{{color:var(--teal);font-weight:800}}
    .callout{{margin-top:22px;border:1px solid #cfe5ee;background:linear-gradient(180deg,var(--blue-soft),#fff);border-radius:22px;padding:22px 24px;color:var(--teal);font-weight:600}}
    .quote-card{{background:var(--teal);color:#e9f5f8;border-radius:28px;padding:28px;box-shadow:var(--shadow)}}
    .quote-card p{{font-size:18px;line-height:1.75}}

    .gallery{{border-top:1px solid var(--line);padding:74px 0}}
    .gallery-head{{max-width:760px;margin-bottom:28px}}
    .gallery-grid{{display:grid;grid-template-columns:1.2fr .8fr;gap:20px}}
    .gallery-grid figure{{position:relative;border-radius:28px;overflow:hidden;box-shadow:var(--shadow);min-height:300px;background:#dfe8ec}}
    .gallery-grid img{{width:100%;height:100%;object-fit:cover;min-height:300px}}
    .gallery-grid figure:nth-child(1){{grid-row:span 2;min-height:620px}}
    .gallery-grid figure:nth-child(1) img{{min-height:620px}}
    figcaption{{position:absolute;right:18px;bottom:18px;color:#fff;background:rgba(7,30,38,.72);backdrop-filter:blur(8px);border:1px solid rgba(255,255,255,.2);padding:8px 12px;border-radius:999px;font-size:14px;font-weight:700}}

    .pdf-section{{border-top:1px solid var(--line);padding:74px 0}}
    .pdf-card{{display:grid;grid-template-columns:minmax(0,1fr) 390px;gap:28px;align-items:start;background:linear-gradient(180deg,#fff,var(--tint));border:1px solid var(--line);border-radius:30px;padding:28px;box-shadow:var(--shadow)}}
    .pdf-viewer{{min-height:620px;border-radius:20px;overflow:hidden;border:1px solid var(--line);background:#fff}}
    .pdf-viewer iframe{{width:100%;height:620px;border:0;display:block;background:#fff}}
    .pdf-info{{padding:8px}}
    .pdf-info h2{{margin-bottom:16px}}
    .pdf-info p{{margin-bottom:18px}}
    .pdf-actions{{display:flex;flex-wrap:wrap;gap:12px;margin-top:22px}}
    .pdf-actions .btn{{padding:12px 17px}}
    .pdf-actions .dark{{background:var(--teal);color:#fff}}
    .pdf-actions .light{{background:#fff;border:1px solid var(--line);color:var(--teal)}}

    .cta{{margin-top:22px;background:linear-gradient(135deg,var(--teal),#0f2c36);color:#fff}}
    .cta-inner{{max-width:var(--max);margin:auto;padding:60px 24px;text-align:center}}
    .cta h2{{color:#fff;max-width:760px;margin:0 auto 16px}}
    .cta h2::after{{right:50%;transform:translateX(50%);background:var(--blue)}}
    .cta p{{max-width:650px;margin:0 auto 24px;color:#d0e5eb;font-size:18px}}
    .wa-btn{{background:#2a8c5e;color:#fff;box-shadow:0 18px 38px -24px #000}}
    .wa-btn:hover{{background:#247a52}}
    .disclaimer{{max-width:var(--max);margin:auto;padding:34px 24px;color:var(--muted);font-size:13.5px;line-height:1.7}}
    .disclaimer h3{{font-size:15px;color:var(--body);font-weight:800;margin-bottom:8px}}
    footer{{background:#0d2d37;color:#a8c6cf;text-align:center;font-size:13.5px;padding:26px 20px;line-height:1.8}}
    footer b{{color:#fff}}
    .floating-wa{{position:fixed;left:22px;bottom:22px;z-index:70;width:56px;height:56px;border-radius:999px;background:#2a8c5e;display:flex;align-items:center;justify-content:center;box-shadow:0 18px 40px -22px #000;text-decoration:none}}
    .floating-wa svg{{width:28px;height:28px}}

    @media (max-width: 960px) {{
      .topbar .inner{{padding:8px 18px}}
      nav a:not(.nav-cta){{display:none}}
      .brand-logo{{width:190px;height:52px}}
      .hero-inner{{grid-template-columns:1fr;padding-top:60px;gap:24px}}
      .hero::after{{background:linear-gradient(180deg,rgba(6,25,31,.35),rgba(6,25,31,.94))}}
      .hero-panel{{max-width:460px}}
      .intro,.content-section,.content-section.reverse,.pdf-card{{grid-template-columns:1fr}}
      .content-section.reverse .section-copy,.content-section.reverse .section-media{{grid-column:auto;grid-row:auto}}
      .aside-card{{position:relative;top:auto}}
      .feature-list.two-col{{grid-template-columns:1fr}}
      .gallery-grid{{grid-template-columns:1fr}}
      .gallery-grid figure:nth-child(1),.gallery-grid figure:nth-child(1) img{{min-height:360px}}
      .pdf-viewer,.pdf-viewer iframe{{min-height:520px;height:520px}}
    }}
    @media (max-width: 620px) {{
      body{{font-size:16.5px;line-height:1.74}}
      .hero-inner,main{{padding-inline:18px}}
      .hero-card{{padding:24px;border-radius:24px}}
      .hero .summary{{font-size:17.5px}}
      .hero-panel{{padding:18px;border-radius:24px}}
      .hero-panel img.logo-large{{height:72px}}
      .content-section,.gallery,.pdf-section{{padding:52px 0}}
      .section-media,.section-media img{{min-height:260px}}
      .pdf-card{{padding:18px;border-radius:24px}}
      .pdf-viewer,.pdf-viewer iframe{{min-height:430px;height:430px}}
      .floating-wa{{left:16px;bottom:16px;width:52px;height:52px}}
    }}
    @media print {{
      .topbar,.progress,.floating-wa,.hero-actions,.pdf-viewer,.cta{{display:none}}
      .hero{{min-height:auto;background:#fff;color:#000}}
      .hero::before,.hero::after{{display:none}}
      .hero-card,.hero-panel{{box-shadow:none;background:#fff;color:#000;border:1px solid #ccc}}
    }}
  </style>
</head>
<body>
  <div class="progress" id="progress"></div>
  <header class="topbar">
    <div class="inner">
      <a class="brand" href="#top" aria-label="Aspect Boutique Investments">
        <img class="brand-logo" src="assets/Horizontal.png" alt="Aspect Boutique Investments">
      </a>
      <nav aria-label="ניווט ראשי">
        <a href="#model">המודל</a>
        <a href="#houston">יוסטון</a>
        <a href="#process">התהליך</a>
        <a href="#proforma">פרופורמה</a>
        <a class="nav-cta" href="{whatsapp_href}" target="_blank" rel="noopener">דברו איתנו</a>
      </nav>
    </div>
  </header>

  <section class="hero" id="top" aria-label="פתיח">
    <div class="hero-inner">
      <div class="hero-card">
        <div class="hero-badge">Aspect Boutique Investments</div>
        <h1>{escape(paras[1])}</h1>
        <div class="subtitle">{escape(paras[2])}</div>
        <p class="summary">{escape(paras[13])}</p>
        <div class="hero-actions">
          <a class="btn primary" href="#proforma">צפייה בדוגמת פרופורמה</a>
          <a class="btn secondary" href="{whatsapp_href}" target="_blank" rel="noopener">שיחה ב־WhatsApp</a>
        </div>
      </div>
      <aside class="hero-panel" aria-label="עיקרי המודל">
        <div class="logo-box"><img class="logo-large" src="assets/Horizontal.png" alt="Aspect Boutique Investments"></div>
        <div class="metrics">
          <div class="metric"><strong>15+ שנות ניסיון</strong><span>ליווי, ניהול וייזום פרויקטים בנדל״ן</span></div>
          <div class="metric"><strong>יוסטון, טקסס</strong><span>שכונות מתפתחות עם ביקוש לשכירות</span></div>
          <div class="metric"><strong>נכס עצמאי</strong><span>בית פרטי בבעלות ישירה של המשקיע</span></div>
        </div>
      </aside>
    </div>
  </section>

  <main>
    <section class="intro" id="model">
      <div class="intro-copy">
        <div class="eyebrow">הבסיס למודל</div>
        {intro}
      </div>
      <aside class="aside-card">
        <h3>תמונה מלאה, לא רק עסקה</h3>
        <p>האתר בנוי סביב אותו קו מנחה שמופיע במסמך: בחינה של כל שרשרת הערך — מחיר הרכישה, התכנון, הביצוע, המימון, הניהול ופוטנציאל התשואה.</p>
        <div class="mini">
          <div>כניסה לשלב הקרקע והתכנון</div>
          <div>התאמת הבית מראש לשוק השכירות</div>
          <div>החזקה ארוכת טווח בנכס עצמאי</div>
        </div>
      </aside>
    </section>

    {section('learned','דרך העבודה','מה למדנו בדרך', learn, 'neighborhood-street.jpg', 'שכונה מתפתחת עם בתים פרטיים')}
    {section('houston','בחירת השוק','מה מצאנו ביוסטון', houston, 'houston-skyline.jpg', 'קו הרקיע של יוסטון', 'reverse')}
    {section('why-houston','מגמות ארוכות טווח','אז למה יוסטון?', why_houston, None, '', 'wide')}
    {section('why-build','היגיון המוצר','למה בחרנו ליזום ולבנות בתים חדשים?', why_build, 'homes-exterior.jpg', 'בתים חדשים ביוסטון')}
    {section('why-now','נקודת הכניסה','למה דווקא עכשיו?', why_now, None, '', 'wide')}
    {section('private-asset','מבנה ההחזקה','להנות מיתרונות הייזום, ולהישאר עם נכס פרטי.', private_asset, None, '', 'wide')}
    {section('role','המעורבות שלנו','התפקיד שלנו בתהליך.', role, None, '', 'wide')}
    {section('capital','מבנה העסקה','תזמון ההון ומנגנון ההנחה.', capital, None, '', 'wide')}
    {section('time','תפיסת הזמן','הזמן כחלק מאסטרטגיה.', time_strategy, None, '', 'wide')}
    {section('process','ניהול הפרויקט','ליווי וניהול מקצה לקצה.', end_to_end, None, '', 'wide')}

    <section class="gallery" aria-label="תמונות מהמסמך">
      <div class="gallery-head">
        <div class="eyebrow">תמונות מתוך חומרי העבודה</div>
        <h2>יוסטון, שכונות מתפתחות ובתים חדשים</h2>
        <p>התמונות מוטמעות כקבצי אתר רגילים, כך שניתן להעלות אותן ל־GitHub יחד עם שאר נכסי האתר ולהחליף אותן בעתיד לפי הצורך.</p>
      </div>
      <div class="gallery-grid">
        <figure><img src="assets/houston-skyline.jpg" alt="קו הרקיע של יוסטון" loading="lazy"><figcaption>יוסטון</figcaption></figure>
        <figure><img src="assets/neighborhood-street.jpg" alt="שכונת בתים פרטיים" loading="lazy"><figcaption>שכונת מגורים</figcaption></figure>
        <figure><img src="assets/homes-exterior.jpg" alt="חזית בתים חדשים" loading="lazy"><figcaption>בתים חדשים</figcaption></figure>
      </div>
    </section>

    {section('handoff','אחרי הבנייה','מהמסירה להחזקה ארוכת טווח.', long_hold, None, '', 'wide')}
    {section('next','שלבי המשך','איך ממשיכים מכאן.', next_steps + '<div class="callout">' + callout_html + '</div>', None, '', 'wide')}

    <section class="pdf-section" id="proforma">
      <div class="pdf-card">
        <div class="pdf-viewer" aria-label="דוגמת פרופורמה מוטמעת">
          <iframe src="assets/proforma-example.pdf#toolbar=1&navpanes=0" title="דוגמת פרופורמה"></iframe>
        </div>
        <div class="pdf-info">
          <div class="eyebrow">דוגמה מספרית</div>
          <h2>דוגמת פרופורמה מלאה</h2>
          <p>קובץ הפרופורמה נשאר כ־PDF עצמאי ונגיש מהאתר. הוא אינו מפורק לטקסט בדף, כדי לשמור על מבנה המסמך, הטבלאות והעיצוב המקוריים.</p>
          <p>במקום קישור למחשבון חיצוני, הדף מציג קישור ישיר לדוגמת הפרופורמה, כך שהמשקיע יכול לפתוח את המסמך המלא בדפדפן או להוריד אותו.</p>
          <div class="pdf-actions">
            <a class="btn dark" href="assets/proforma-example.pdf" target="_blank" rel="noopener">פתיחת ה־PDF</a>
            <a class="btn light" href="assets/proforma-example.pdf" download>הורדת ה־PDF</a>
          </div>
        </div>
      </div>
    </section>
  </main>

  <section class="cta" id="contact">
    <div class="cta-inner">
      <h2>רוצים להבין אם המודל מתאים לכם?</h2>
      <p>השיחה הראשונה נועדה להבין את מסגרת ההשקעה, טווח הזמן, רמת הסיכון המתאימה והאם נכון להתקדם לבחינת פרויקט קונקרטי.</p>
      <a class="btn wa-btn" href="{whatsapp_href}" target="_blank" rel="noopener" aria-label="פתיחת שיחה בוואטסאפ">
        <svg viewBox="0 0 24 24" width="21" height="21" fill="#fff" xmlns="http://www.w3.org/2000/svg" aria-hidden="true"><path d="M12.04 2c-5.46 0-9.91 4.45-9.91 9.91 0 1.75.46 3.45 1.32 4.95L2.05 22l5.25-1.38c1.45.79 3.08 1.21 4.74 1.21h.01c5.46 0 9.91-4.45 9.91-9.91 0-2.65-1.03-5.14-2.9-7.01A9.82 9.82 0 0 0 12.04 2zm5.8 14.13c-.24.68-1.42 1.31-1.95 1.36-.5.05-.96.24-3.23-.67-2.73-1.08-4.46-3.86-4.6-4.04-.13-.18-1.1-1.46-1.1-2.79 0-1.33.7-1.98.95-2.25.24-.27.53-.34.71-.34.18 0 .35 0 .51.01.16.01.39-.06.6.46.24.58.81 2 .88 2.14.07.14.12.31.02.49-.09.18-.14.29-.27.45-.14.16-.29.36-.41.48-.14.14-.28.29-.12.56.16.27.71 1.17 1.53 1.9 1.05.94 1.94 1.23 2.21 1.37.27.14.43.12.59-.07.16-.18.68-.79.86-1.06.18-.27.36-.22.6-.13.24.09 1.55.73 1.82.86.27.14.45.2.51.31.07.12.07.66-.17 1.32z"/></svg>
        פתיחת שיחה ב־WhatsApp
      </a>
    </div>
  </section>

  <div class="disclaimer">
    <h3>{escape(paras[93])}</h3>
    <p>{escape(paras[94])}</p>
  </div>

  <footer>
    <b>Aspect Boutique Investments</b><br>
    המסמך מיועד לנמענים רלוונטיים בלבד ואינו מיועד להפצה פומבית או שימוש מסחרי ללא אישור מראש ובכתב.
  </footer>

  <a class="floating-wa" href="{whatsapp_href}" target="_blank" rel="noopener" aria-label="WhatsApp">
    <svg viewBox="0 0 24 24" fill="#fff" xmlns="http://www.w3.org/2000/svg" aria-hidden="true"><path d="M12.04 2c-5.46 0-9.91 4.45-9.91 9.91 0 1.75.46 3.45 1.32 4.95L2.05 22l5.25-1.38c1.45.79 3.08 1.21 4.74 1.21h.01c5.46 0 9.91-4.45 9.91-9.91 0-2.65-1.03-5.14-2.9-7.01A9.82 9.82 0 0 0 12.04 2zm5.8 14.13c-.24.68-1.42 1.31-1.95 1.36-.5.05-.96.24-3.23-.67-2.73-1.08-4.46-3.86-4.6-4.04-.13-.18-1.1-1.46-1.1-2.79 0-1.33.7-1.98.95-2.25.24-.27.53-.34.71-.34.18 0 .35 0 .51.01.16.01.39-.06.6.46.24.58.81 2 .88 2.14.07.14.12.31.02.49-.09.18-.14.29-.27.45-.14.16-.29.36-.41.48-.14.14-.28.29-.12.56.16.27.71 1.17 1.53 1.9 1.05.94 1.94 1.23 2.21 1.37.27.14.43.12.59-.07.16-.18.68-.79.86-1.06.18-.27.36-.22.6-.13.24.09 1.55.73 1.82.86.27.14.45.2.51.31.07.12.07.66-.17 1.32z"/></svg>
  </a>

  <script>
    const progress = document.getElementById('progress');
    function updateProgress() {{
      const page = document.documentElement;
      const max = page.scrollHeight - page.clientHeight || 1;
      progress.style.width = `${{(page.scrollTop / max) * 100}}%`;
    }}
    addEventListener('scroll', updateProgress, {{ passive: true }});
    updateProgress();
  </script>
</body>
</html>
'''
(SITE / 'index.html').write_text(html, encoding='utf-8')
print(SITE / 'index.html')
